from docx import Document

# Load the Word document
file_path = r"C:\Users\v venkata ramana\OneDrive\Desktop\python_project\vvr.docx"
doc = Document(file_path)

# Read all paragraphs
for para in doc.paragraphs:
    print(para.text)

# Modify the first paragraph
if doc.paragraphs:
    doc.paragraphs[0].text = "This is the updated first paragraph."

# Add a new paragraph
doc.add_paragraph("This is a new paragraph added by python-docx.")

# Save the changes to a new file
doc.save(r"C:\Users\v venkata ramana\OneDrive\Desktop\python_project\vvr.docx")